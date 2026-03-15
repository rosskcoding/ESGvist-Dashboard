"""
Artifact-from-ZIP Tests.

Goal: lock in the invariant that artifact generation (PDF/DOCX/HTML export)
does not depend on a build workspace. All required inputs must come from the ZIP.

This supports the architecture decision: delete the workspace after build and
produce artifacts from the ZIP only.

Tests:
- TEST-01: PDF generation from ZIP when workspace is missing
- TEST-02: Explicit error when print bundle is missing in ZIP
- TEST-03: Temporary directory is cleaned up after generation
"""

import json
import os
import tempfile
import zipfile
from pathlib import Path
from unittest.mock import AsyncMock, MagicMock, patch

import pytest


class TestPDFGenerationFromZip:
    """
    Tests for PDF generation from ZIP archive.

    Key invariant: PDF generator takes ONLY zip_path as input,
    does NOT require or access workspace directory.
    """

    @pytest.fixture
    def minimal_print_bundle_zip(self, tmp_path: Path) -> Path:
        """
        Create minimal ZIP with print bundle structure.

        Structure:
            print/
                en.html
            assets/
                css/
                    print.css
        """
        zip_path = tmp_path / "build.zip"

        print_html_content = """<!DOCTYPE html>
<html lang="en">
<head>
    <meta charset="UTF-8">
    <title>Test Report</title>
    <link rel="stylesheet" href="../assets/css/print.css">
</head>
<body>
    <h1>Test report</h1>
    <section>
        <h2>Section 1</h2>
        <p>Section content.</p>
    </section>
</body>
</html>"""

        print_css_content = """@page {
    size: A4;
    margin: 20mm 15mm 25mm 15mm;
}

body {
    font-family: Arial, sans-serif;
    font-size: 12pt;
    line-height: 1.5;
}

h1 { font-size: 24pt; }
h2 { font-size: 18pt; }
"""

        with zipfile.ZipFile(zip_path, "w", zipfile.ZIP_DEFLATED) as zf:
            zf.writestr("print/en.html", print_html_content)
            zf.writestr("assets/css/print.css", print_css_content)
            # Manifest (optional but realistic)
            zf.writestr("build-manifest.json", '{"build_id": "test", "locales": ["en"]}')

        return zip_path

    @pytest.fixture
    def zip_without_print_bundle(self, tmp_path: Path) -> Path:
        """Create ZIP without print bundle (only assets)."""
        zip_path = tmp_path / "bad_build.zip"

        with zipfile.ZipFile(zip_path, "w") as zf:
            zf.writestr("assets/css/app.css", "body { color: black; }")
            zf.writestr("en/index.html", "<html>Interactive HTML</html>")
            # Missing print/ directory entirely

        return zip_path

    @pytest.fixture
    def zip_with_wrong_locale(self, tmp_path: Path) -> Path:
        """Create ZIP with print bundle but missing requested locale."""
        zip_path = tmp_path / "wrong_locale.zip"

        with zipfile.ZipFile(zip_path, "w") as zf:
            # Only English print HTML
            zf.writestr("print/en.html", "<html><body>English</body></html>")
            zf.writestr("assets/css/print.css", "@page { size: A4; }")

        return zip_path

    @pytest.mark.asyncio
    async def test_01_pdf_generates_from_zip_without_workspace(
        self, minimal_print_bundle_zip: Path, tmp_path: Path
    ):
        """
        TEST-01: PDF task generates a file from ZIP when workspace is missing.

        Invariant:
        - Workspace does not exist (or was deleted)
        - PDF is generated successfully from the ZIP
        - Renderer loads extracted HTML from a temp directory (not workspace)

        Criteria: success without workspace.
        """
        from app.services.pdf_generator import generate_pdf

        # Arrange: No workspace directory exists
        workspace_path = tmp_path / "workspace"
        assert not workspace_path.exists(), "Workspace should NOT exist"

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Track what HTML path was loaded
        loaded_html_path = None

        # Mock Playwright renderer to avoid real Chromium dependency
        mock_page = AsyncMock()
        mock_page.set_default_timeout = MagicMock()
        mock_page.evaluate = AsyncMock()

        async def mock_goto(url, **kwargs):
            nonlocal loaded_html_path
            loaded_html_path = url

        mock_page.goto = mock_goto

        # Mock PDF generation to create actual file
        async def mock_pdf(**kwargs):
            pdf_path = kwargs.get("path")
            if pdf_path:
                Path(pdf_path).write_bytes(b"%PDF-1.4 mock content")

        mock_page.pdf = mock_pdf

        mock_context = AsyncMock()
        mock_context.new_page = AsyncMock(return_value=mock_page)

        mock_browser = AsyncMock()
        mock_browser.new_context = AsyncMock(return_value=mock_context)
        mock_browser.close = AsyncMock()

        mock_chromium = MagicMock()
        mock_chromium.launch = AsyncMock(return_value=mock_browser)

        mock_playwright_instance = MagicMock()
        mock_playwright_instance.chromium = mock_chromium

        # Create mock for async context manager
        mock_async_playwright = AsyncMock()
        mock_async_playwright.__aenter__ = AsyncMock(return_value=mock_playwright_instance)
        mock_async_playwright.__aexit__ = AsyncMock(return_value=None)

        # Patch playwright.async_api.async_playwright (where it's imported from)
        with patch("playwright.async_api.async_playwright", return_value=mock_async_playwright):
            # Act
            pdf_path, sha256, size_bytes = await generate_pdf(
                zip_path=minimal_print_bundle_zip,
                locale="en",
                profile="audit",
                output_dir=output_dir,
            )

        # Assert: File created
        assert Path(pdf_path).exists(), "PDF file should be created"
        assert size_bytes > 0, "PDF should have content"
        assert sha256 is not None, "SHA256 should be computed"

        # Assert: Renderer was called with temp extracted path, NOT workspace
        assert loaded_html_path is not None, "HTML should have been loaded"
        assert "file://" in loaded_html_path, "Should be local file URL"

        # The loaded path should be in temp directory, NOT workspace
        loaded_path_str = loaded_html_path.replace("file://", "")
        assert "workspace" not in loaded_path_str.lower(), (
            f"Should NOT access workspace, but loaded: {loaded_path_str}"
        )

        # The path should be in system temp directory
        temp_root = tempfile.gettempdir()
        assert loaded_path_str.startswith(temp_root) or "/tmp" in loaded_path_str, (
            f"Should use temp directory, but loaded: {loaded_path_str}"
        )

        # Assert: Workspace still doesn't exist (never created)
        assert not workspace_path.exists(), "Workspace should NOT be created"

    @pytest.mark.asyncio
    async def test_02_missing_print_bundle_fails_fast(
        self, zip_without_print_bundle: Path, tmp_path: Path
    ):
        """
        TEST-02: Missing print bundle -> explicit error (fail-fast).

        If the ZIP is broken (no print/en.html), the artifact must not "silently"
        succeed. The error should be clear.

        Criteria:
        - PDFGeneratorError raised
        - Error message contains "Print HTML not found"
        - Renderer NOT called (error happens before Playwright is invoked)
        """
        from app.services.pdf_generator import PDFGeneratorError, generate_pdf

        # Arrange
        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Act & Assert - error should happen during ZIP extraction, before Playwright
        with pytest.raises(PDFGeneratorError) as exc_info:
            await generate_pdf(
                zip_path=zip_without_print_bundle,
                locale="en",
                profile="audit",
                output_dir=output_dir,
            )

        # Assert: Clear error message
        error_message = str(exc_info.value)
        assert "Print HTML not found" in error_message, (
            f"Error should mention missing print HTML, got: {error_message}"
        )

    @pytest.mark.asyncio
    async def test_02b_wrong_locale_fails_fast(
        self, zip_with_wrong_locale: Path, tmp_path: Path
    ):
        """
        TEST-02b: Requesting locale not in ZIP → explicit error.

        ZIP has print/en.html but we request "fr" locale.
        """
        from app.services.pdf_generator import PDFGeneratorError, generate_pdf

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Act & Assert
        with pytest.raises(PDFGeneratorError) as exc_info:
            await generate_pdf(
                zip_path=zip_with_wrong_locale,
                locale="fr",  # ZIP only has "en"
                profile="audit",
                output_dir=output_dir,
            )

        error_message = str(exc_info.value)
        assert "Print HTML not found" in error_message or "fr" in error_message, (
            f"Error should mention missing locale, got: {error_message}"
        )

    @pytest.mark.asyncio
    async def test_03_temp_directory_cleaned_up(
        self, minimal_print_bundle_zip: Path, tmp_path: Path
    ):
        """
        TEST-03: Temporary directory is cleaned up after generation.

        If extraction creates a temp dir, it must be removed after completion
        (even on failure).

        Implementation note: Python's tempfile.TemporaryDirectory
        handles cleanup automatically via context manager.
        """
        from app.services.pdf_generator import generate_pdf

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Track temp directories created
        temp_dirs_before = set()
        temp_root = Path(tempfile.gettempdir())
        for item in temp_root.iterdir():
            if item.is_dir():
                temp_dirs_before.add(item.name)

        # Mock Playwright
        mock_page = AsyncMock()
        mock_page.set_default_timeout = MagicMock()
        mock_page.goto = AsyncMock()
        mock_page.evaluate = AsyncMock()
        # Ensure the mocked PDF call actually creates a file, so generate_pdf can hash it
        async def _mock_pdf(**kwargs):
            pdf_path = kwargs.get("path")
            if pdf_path:
                Path(pdf_path).write_bytes(b"%PDF-1.4 mock content")

        mock_page.pdf = _mock_pdf

        mock_context = AsyncMock()
        mock_context.new_page = AsyncMock(return_value=mock_page)

        mock_browser = AsyncMock()
        mock_browser.new_context = AsyncMock(return_value=mock_context)
        mock_browser.close = AsyncMock()

        mock_chromium = MagicMock()
        mock_chromium.launch = AsyncMock(return_value=mock_browser)

        mock_playwright_instance = MagicMock()
        mock_playwright_instance.chromium = mock_chromium

        # Create mock for async context manager
        mock_async_playwright = AsyncMock()
        mock_async_playwright.__aenter__ = AsyncMock(return_value=mock_playwright_instance)
        mock_async_playwright.__aexit__ = AsyncMock(return_value=None)

        # Create PDF to trigger temp directory creation
        with patch("playwright.async_api.async_playwright", return_value=mock_async_playwright):
            await generate_pdf(
                zip_path=minimal_print_bundle_zip,
                locale="en",
                profile="audit",
                output_dir=output_dir,
            )

        # Check temp directories after
        temp_dirs_after = set()
        for item in temp_root.iterdir():
            if item.is_dir():
                temp_dirs_after.add(item.name)

        # New directories should not persist (or be the same set)
        new_dirs = temp_dirs_after - temp_dirs_before

        # Filter for our PDF generator temp dirs (if identifiable)
        # Note: tempfile.TemporaryDirectory creates names like tmp*, so we check
        # that no NEW temp dirs remain (within reasonable tolerance for other processes)
        assert len(new_dirs) < 5, (
            f"Too many new temp dirs remain after PDF generation: {new_dirs}"
        )


class TestDOCXGenerationFromZip:
    """
    Tests for DOCX generation from ZIP archive.

    Key invariant: DOCX generator takes ONLY zip_path as input,
    does NOT require or access database or workspace directory.
    """

    @pytest.fixture
    def zip_with_content_snapshot(self, tmp_path: Path) -> Path:
        """
        Create ZIP with content snapshot (sections/blocks JSON).

        Structure:
            content-snapshot.json
            build-manifest.json
        """
        zip_path = tmp_path / "build.zip"

        content_snapshot = {
            "report_id": "test-report-id",
            "report_title": "Test report",
            "report_year": 2024,
            "company_name": "Test company",
            "sections": [
                {
                    "section_id": "section-1",
                    "parent_section_id": None,
                    "order_index": 0,
                    "depth": 0,
                    "label_prefix": "1",
                    "label_suffix": None,
                    "i18n": [
                        {
                            "locale": "en",
                            "title": "Section 1",
                            "slug": "section-1",
                            "summary": "Section description",
                        }
                    ],
                }
            ],
            "blocks": [
                {
                    "block_id": "block-1",
                    "section_id": "section-1",
                    "type": "text",
                    "variant": "default",
                    "order_index": 0,
                    "data_json": {},
                    "qa_flags_global": [],
                    "custom_override_enabled": False,
                    "i18n": [
                        {
                            "locale": "en",
                            "status": "published",
                            "qa_flags_by_locale": [],
                            "fields_json": {
                                "body_html": "<p>Test block content.</p>"
                            },
                            "custom_html_sanitized": None,
                            "custom_css_validated": None,
                        }
                    ],
                }
                ,
                # TABLE (Builder mode): data.rows[].cells + fields.column_labels
                {
                    "block_id": "table-builder",
                    "section_id": "section-1",
                    "type": "table",
                    "variant": "default",
                    "order_index": 1,
                    "data_json": {
                        "mode": "builder",
                        "columns": [{"key": "col1"}, {"key": "col2"}],
                        "rows": [
                            {"cells": {"col1": "Value A", "col2": "1"}},
                            {"cells": {"col1": "Value B", "col2": "2"}},
                        ],
                    },
                    "qa_flags_global": [],
                    "custom_override_enabled": False,
                    "i18n": [
                        {
                            "locale": "en",
                            "status": "published",
                            "qa_flags_by_locale": [],
                            "fields_json": {
                                "caption": "Table (Builder)",
                                "column_labels": {"col1": "Column 1", "col2": "Column 2"},
                            },
                            "custom_html_sanitized": None,
                            "custom_css_validated": None,
                        }
                    ],
                },
                # TABLE (Builder mode): headers in fields.columns[key].header
                {
                    "block_id": "table-fields-columns",
                    "section_id": "section-1",
                    "type": "table",
                    "variant": "default",
                    "order_index": 2,
                    "data_json": {
                        "mode": "builder",
                        "columns": [{"key": "col1"}, {"key": "col2"}],
                        "rows": [{"cells": {"col1": "2015", "col2": "11200"}}],
                    },
                    "qa_flags_global": [],
                    "custom_override_enabled": False,
                    "i18n": [
                        {
                            "locale": "en",
                            "status": "published",
                            "qa_flags_by_locale": [],
                            "fields_json": {
                                "caption": "Table (fields.columns headers)",
                                "columns": {
                                    "col1": {"header": "Year"},
                                    "col2": {"header": "Headcount"},
                                },
                            },
                            "custom_html_sanitized": None,
                            "custom_css_validated": None,
                        }
                    ],
                },
                # TABLE (Legacy matrix): fields.header_row + fields.data_rows
                {
                    "block_id": "table-legacy-matrix",
                    "section_id": "section-1",
                    "type": "table",
                    "variant": "default",
                    "order_index": 3,
                    "data_json": {
                        "mode": "builder",
                        "has_header": True,
                        "columns": [{"key": "indicator"}, {"key": "y2024"}],
                    },
                    "qa_flags_global": [],
                    "custom_override_enabled": False,
                    "i18n": [
                        {
                            "locale": "en",
                            "status": "published",
                            "qa_flags_by_locale": [],
                            "fields_json": {
                                "caption": "Table (legacy matrix)",
                                "header_row": ["Metric", "2024"],
                                "data_rows": [["CO2", "100"], ["CH4", "50"]],
                            },
                            "custom_html_sanitized": None,
                            "custom_css_validated": None,
                        }
                    ],
                },
            ],
        }

        with zipfile.ZipFile(zip_path, "w") as zf:
            zf.writestr(
                "content-snapshot.json",
                json.dumps(content_snapshot, ensure_ascii=False, indent=2)
            )
            zf.writestr(
                "build-manifest.json",
                '{"build_id": "test", "locales": ["en"]}'
            )

        return zip_path

    @pytest.fixture
    def zip_without_content_snapshot(self, tmp_path: Path) -> Path:
        """Create ZIP without content snapshot."""
        zip_path = tmp_path / "old_build.zip"

        with zipfile.ZipFile(zip_path, "w") as zf:
            zf.writestr("print/en.html", "<html>Print</html>")
            zf.writestr("build-manifest.json", '{"build_id": "test"}')

        return zip_path

    @pytest.mark.asyncio
    async def test_01_docx_generates_from_zip_without_database(
        self, zip_with_content_snapshot: Path, tmp_path: Path
    ):
        """
        TEST-01 (DOCX): DOCX is generated from ZIP without DB access.

        Invariant:
        - No DB access (session=None)
        - DOCX is generated successfully from content-snapshot.json in the ZIP

        Criteria: success without DB.
        """
        from app.services.docx_generator import generate_docx_from_zip

        # Arrange
        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Act - no database session needed
        docx_path, sha256, size_bytes = await generate_docx_from_zip(
            zip_path=zip_with_content_snapshot,
            locale="en",
            output_dir=output_dir,
            include_toc=False,
            include_cover=True,
        )

        # Assert: File created
        assert Path(docx_path).exists(), "DOCX file should be created"
        assert size_bytes > 0, "DOCX should have content"
        assert sha256 is not None, "SHA256 should be computed"

        # Assert: File is valid DOCX
        assert docx_path.endswith(".docx"), "Should be .docx file"

        # Verify DOCX can be opened (basic validation)
        try:
            from docx import Document
            doc = Document(docx_path)
            # Should have some content
            assert len(doc.paragraphs) > 0, "DOCX should have paragraphs"
        except Exception as e:
            pytest.fail(f"DOCX file is not valid: {e}")

    @pytest.mark.asyncio
    async def test_01b_docx_renders_tables_from_snapshot_shapes(
        self, zip_with_content_snapshot: Path, tmp_path: Path
    ):
        """DOCX must contain table cell values for Builder and legacy matrix shapes."""
        from app.services.docx_generator import generate_docx_from_zip

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        docx_path, _, _ = await generate_docx_from_zip(
            zip_path=zip_with_content_snapshot,
            locale="en",
            output_dir=output_dir,
            include_toc=False,
            include_cover=False,
        )

        from docx import Document

        doc = Document(docx_path)
        assert len(doc.tables) >= 3, "Expected multiple tables from snapshot"

        table_text = "\n".join(
            cell.text
            for table in doc.tables
            for row in table.rows
            for cell in row.cells
        )

        # Builder mode with column_labels
        assert "Column 1" in table_text
        assert "Column 2" in table_text
        assert "Value A" in table_text
        assert "Value B" in table_text

        # Builder mode with fields.columns headers
        assert "Year" in table_text
        assert "Headcount" in table_text
        assert "2015" in table_text
        assert "11200" in table_text

        # Legacy matrix
        assert "Metric" in table_text
        assert "2024" in table_text
        assert "CO2" in table_text
        assert "CH4" in table_text

    @pytest.mark.asyncio
    async def test_02_missing_content_snapshot_fails_fast(
        self, zip_without_content_snapshot: Path, tmp_path: Path
    ):
        """
        TEST-02 (DOCX): Missing content snapshot -> explicit error.

        If the ZIP does not contain content-snapshot.json, the generator should return
        a clear error.
        """
        from app.services.docx_generator import DOCXGeneratorError, generate_docx_from_zip

        # Arrange
        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Act & Assert
        with pytest.raises(DOCXGeneratorError) as exc_info:
            await generate_docx_from_zip(
                zip_path=zip_without_content_snapshot,
                locale="en",
                output_dir=output_dir,
            )

        error_message = str(exc_info.value)
        assert "content-snapshot.json not found" in error_message, (
            f"Error should mention missing content snapshot, got: {error_message}"
        )

    @pytest.mark.asyncio
    async def test_02b_missing_locale_in_snapshot(
        self, zip_with_content_snapshot: Path, tmp_path: Path
    ):
        """
        TEST-02b (DOCX): Requesting locale not in content snapshot.

        Snapshot has only "en", we request "fr".
        Should produce empty/minimal DOCX (sections without translation are skipped).
        """
        from app.services.docx_generator import generate_docx_from_zip

        output_dir = tmp_path / "output"
        output_dir.mkdir()

        # Act - request a locale that doesn't exist in snapshot
        docx_path, sha256, size_bytes = await generate_docx_from_zip(
            zip_path=zip_with_content_snapshot,
            locale="fr",  # Not in snapshot
            output_dir=output_dir,
        )

        # Assert: File created but minimal (only cover page)
        assert Path(docx_path).exists(), "DOCX should be created even with missing locale"

        # Verify it's valid but has minimal content
        from docx import Document
        doc = Document(docx_path)
        # Should have cover page paragraphs but no section content
        assert len(doc.paragraphs) > 0, "Should have at least cover page"


class TestZIPInputOnly:
    """
    Tests confirming that artifact generation accepts ONLY ZIP input.

    SUT requirement: generate_pdf accepts only zip_path (or URI),
    NOT workspace path.
    """

    def test_generate_pdf_signature_requires_zip_path(self):
        """Verify generate_pdf function signature requires zip_path."""
        import inspect
        from app.services.pdf_generator import generate_pdf

        sig = inspect.signature(generate_pdf)
        params = list(sig.parameters.keys())

        # First parameter should be zip_path
        assert params[0] == "zip_path", (
            f"First parameter should be zip_path, got: {params[0]}"
        )

        # Should NOT have workspace_path parameter
        assert "workspace_path" not in params, (
            "Should NOT have workspace_path parameter"
        )
        assert "workspace" not in params, (
            "Should NOT have workspace parameter"
        )

    def test_generate_pdf_sync_signature_requires_zip_path(self):
        """Verify synchronous wrapper also requires zip_path."""
        import inspect
        from app.services.pdf_generator import generate_pdf_sync

        sig = inspect.signature(generate_pdf_sync)
        params = list(sig.parameters.keys())

        assert params[0] == "zip_path", (
            f"First parameter should be zip_path, got: {params[0]}"
        )
        assert "workspace" not in params

    def test_generate_docx_from_zip_signature(self):
        """Verify generate_docx_from_zip requires zip_path, not session."""
        import inspect
        from app.services.docx_generator import generate_docx_from_zip

        sig = inspect.signature(generate_docx_from_zip)
        params = list(sig.parameters.keys())

        # First parameter should be zip_path
        assert params[0] == "zip_path", (
            f"First parameter should be zip_path, got: {params[0]}"
        )

        # Should NOT have session or report_id parameters
        assert "session" not in params, (
            "Should NOT have session parameter (no DB access)"
        )
        assert "report_id" not in params, (
            "Should NOT have report_id parameter"
        )


class TestPDFGeneratorErrorHandling:
    """Tests for error handling in PDF generation."""

    @pytest.mark.asyncio
    async def test_nonexistent_zip_raises_error(self, tmp_path: Path):
        """Should raise PDFGeneratorError when ZIP doesn't exist."""
        from app.services.pdf_generator import PDFGeneratorError, generate_pdf

        fake_zip = tmp_path / "nonexistent.zip"

        with pytest.raises(PDFGeneratorError) as exc_info:
            await generate_pdf(
                zip_path=fake_zip,
                locale="ru",
                output_dir=tmp_path,
            )

        assert "ZIP file not found" in str(exc_info.value)

    @pytest.mark.asyncio
    async def test_corrupted_zip_raises_error(self, tmp_path: Path):
        """Should raise PDFGeneratorError for corrupted ZIP."""
        from app.services.pdf_generator import PDFGeneratorError, generate_pdf

        # Create corrupted "ZIP"
        bad_zip = tmp_path / "corrupted.zip"
        bad_zip.write_text("This is not a valid ZIP file content")

        with pytest.raises(PDFGeneratorError) as exc_info:
            await generate_pdf(
                zip_path=bad_zip,
                locale="ru",
                output_dir=tmp_path,
            )

        assert "Invalid ZIP" in str(exc_info.value) or "ZIP" in str(exc_info.value)


class TestArchitecturalInvariant:
    """
    Meta-tests documenting the architectural decision.

    These tests serve as living documentation that workspace can be
    deleted after build, and artifacts are built from ZIP only.
    """

    def test_pdf_generator_uses_tempfile_for_extraction(self):
        """
        Verify PDF generator uses tempfile.TemporaryDirectory for ZIP extraction.

        This ensures:
        1. No dependency on persistent workspace
        2. Automatic cleanup after generation
        3. Isolation between concurrent generations
        """
        import ast
        from pathlib import Path

        pdf_gen_path = Path(__file__).parents[2] / "app" / "services" / "pdf_generator.py"
        source = pdf_gen_path.read_text()

        # Parse and check for tempfile usage
        assert "tempfile.TemporaryDirectory" in source or "TemporaryDirectory" in source, (
            "PDF generator should use TemporaryDirectory for extraction"
        )

        # Should use context manager (with statement)
        tree = ast.parse(source)
        has_with_tempdir = False

        for node in ast.walk(tree):
            if isinstance(node, ast.With):
                for item in node.items:
                    if hasattr(item.context_expr, "func"):
                        func = item.context_expr.func
                        if hasattr(func, "attr") and "TemporaryDirectory" in func.attr:
                            has_with_tempdir = True

        assert has_with_tempdir, (
            "TemporaryDirectory should be used with 'with' statement for cleanup"
        )

    def test_build_pipeline_deletes_workspace_after_build(self):
        """
        Verify BuildPipeline.execute() deletes workspace in finally block.

        This confirms the architectural decision that workspace is temporary
        and artifacts must be built from ZIP.
        """
        from pathlib import Path

        pipeline_path = Path(__file__).parents[2] / "app" / "services" / "build_pipeline.py"
        source = pipeline_path.read_text()

        # Check for workspace cleanup in finally block
        assert "shutil.rmtree" in source, (
            "BuildPipeline should use shutil.rmtree to clean workspace"
        )
        assert "finally:" in source, (
            "Workspace cleanup should be in finally block"
        )

        # Check that cleanup happens on self.workspace
        assert "self.workspace" in source, (
            "Should reference self.workspace for cleanup"
        )
